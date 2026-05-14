"""
Клонирует образец отчёта .docx и заменяет организацию на ООО «Альфасервис».
Сохраняет стили Word: текст собирается из runs, результат пишется в первый run.
"""
from pathlib import Path

from docx import Document

SOURCE = Path(r"c:\Users\nazar\Downloads\4123_obrazets_otcheta_preddip_prak (1).docx")
TARGET = Path(r"c:\Users\nazar\OneDrive\Рабочий стол\diplom-main\Practice_Report_Alfaservice_exact.docx")


# Порядок: более длинные фразы — раньше (важно для подстрок)
_REPLACEMENTS = [
    (
        'ГАПОУ «Перевозский строительный колледж», г. Перевоз, Проспект Советский , д. 27.',
        'ООО «Альфасервис», адрес места прохождения практики: _____________________________.',
    ),
    (
        'Государственное автономное профессиональное образовательное учреждение «Перевозский строительный колледж» (ГАПОУ «Перевозский строительный колледж») является некоммерческой организацией, созданной для оказания услуг в целях обеспечения реализации предусмотренных законодательством Российской Федерации полномочий в сфере образования.',
        'Общество с ограниченной ответственностью «Альфасервис» (ООО «Альфасервис») осуществляет предпринимательскую деятельность в соответствии с уставом и законодательством РФ; основные направления работ (по данным практики): _____________________________.',
    ),
    (
        'Я проходил(а) преддипломную практику в организации: ГАПОУ «Перевозский строительный колледж».',
        'Я проходил(а) преддипломную практику в организации ООО «Альфасервис».',
    ),
    (
        'от ГАПОУ “Перевозский строительный колледж”',
        'от образовательной организации _____________________________',
    ),
    (
        'от ГАПОУ «Перевозский строительный колледж»',
        'от образовательной организации _____________________________',
    ),
    ('ГАПОУ «Перевозский строительный колледж»', 'ООО «Альфасервис»'),
    (
        'Государственное автономное профессиональное образовательное учреждение',
        'Образовательная организация (при оформлении отчёта в учебном заведении)',
    ),
    ('“Перевозский строительный колледж”', '_____________________________'),
    ('«Перевозский строительный колледж»', 'ООО «Альфасервис»'),
    ('Перевозский строительный колледж', 'ООО «Альфасервис»'),
    ('г. Перевоз', '_________________'),
    ('Перевоз 20', '________________ 20'),
    ('установленной в колледже', 'установленной в образовательной организации'),
    ('ООО_SportGoods', 'ООО «Альфасервис»'),
    ('OOO_SportGoods', 'ООО «Альфасервис»'),
    ('OxygenStore', 'ООО «Альфасервис»'),
    ('проектная команда «OxygenStore»', 'ООО «Альфасервис»'),
]

REPLACEMENTS = sorted(_REPLACEMENTS, key=lambda x: len(x[0]), reverse=True)


def apply_replacements(text: str) -> str:
    for old, new in REPLACEMENTS:
        text = text.replace(old, new)
    return text


def replace_in_paragraph(paragraph) -> None:
    if not paragraph.runs:
        return
    text = ''.join(run.text for run in paragraph.runs)
    new_text = apply_replacements(text)
    if new_text == text:
        return
    paragraph.runs[0].text = new_text
    for r in paragraph.runs[1:]:
        r.text = ''


def process_table(table) -> None:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_in_paragraph(paragraph)
            for nested in cell.tables:
                process_table(nested)


def process_document(doc) -> None:
    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph)
    for table in doc.tables:
        process_table(table)


def process_header_footer(part) -> None:
    for paragraph in part.paragraphs:
        replace_in_paragraph(paragraph)
    for table in part.tables:
        process_table(table)


def main() -> None:
    if not SOURCE.exists():
        raise FileNotFoundError(f"Файл-образец не найден: {SOURCE}")

    doc = Document(str(SOURCE))
    process_document(doc)

    for section in doc.sections:
        process_header_footer(section.header)
        process_header_footer(section.footer)

    doc.save(str(TARGET))
    print(TARGET)


if __name__ == "__main__":
    main()
